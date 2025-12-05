"""
Word文書生成サービス
"""
from docx import Document
from docx.shared import Pt
from datetime import datetime
from pathlib import Path
from typing import List
from app.domain.summary import Summary
from app.config import Config

class DocxService:
    """Word文書生成サービス"""
    
    def __init__(self):
        pass
    
    def generate_document(self, summaries: List[Summary], formatted_content: str = None, output_path: Path = None) -> Path:
        """
        議事録Wordファイルを生成
        
        Args:
            summaries: 要約データリスト
            formatted_content: Gemini APIで整形済みのコンテンツ（オプション）
            output_path: 出力ファイルパス（Noneの場合は自動生成）
            
        Returns:
            生成されたファイルのパス
        """
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = Config.OUTPUTS_DIR / f"議事録_{timestamp}.docx"
        
        doc = Document()
        
        # タイトル
        title = doc.add_heading('議事録', level=1)
        
        # 整形済みコンテンツがある場合はそれを使用
        if formatted_content:
            # Markdown形式のコンテンツを解析してWord文書に変換
            lines = formatted_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph('')
                    continue
                
                # 見出し
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                # リスト項目
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                # その他は通常の段落
                else:
                    doc.add_paragraph(line)
        else:
            # 従来の方式（カテゴリごと）
            doc.add_paragraph(f'日付: {datetime.now().strftime("%Y年%m月%d日")}')
            doc.add_paragraph(f'場所: （未入力）')
            doc.add_paragraph('')
            
            # カテゴリごとにグループ化
            categories = {}
            for summary in summaries:
                if summary.category not in categories:
                    categories[summary.category] = []
                categories[summary.category].append(summary)
            
            # セクションごとに出力
            for category, items in categories.items():
                doc.add_heading(f'■ {category}', level=2)
                
                for item in items:
                    # 質問
                    p = doc.add_paragraph()
                    p.add_run(f'{item.question_id}. {item.question_text}').bold = True
                    
                    # 要約
                    doc.add_paragraph(f'要約: {item.summary_text}')
                    doc.add_paragraph('')
        
        # ファイル保存
        doc.save(output_path)
        
        return output_path
